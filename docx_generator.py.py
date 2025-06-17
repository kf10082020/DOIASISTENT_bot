from docx import Document

def generate_docx(metadata: dict, output_path: str):
    doc = Document()
    doc.add_heading(metadata.get("title", "—"), level=1)
    doc.add_paragraph(f'Aaoi?u: {metadata.get("authors", "—")}')
    doc.add_paragraph(f'?o?iae: {metadata.get("journal", "—")} ({metadata.get("issued", "—")})')
    doc.add_paragraph(f'Oii: {metadata.get("volume", "—")}  ?{metadata.get("issue", "—")} no?. {metadata.get("pages", "—")}')

    doc.add_heading("Aiiioaoey", level=2)
    doc.add_paragraph(metadata.get("abstract", "—"))

    doc.add_heading("Auaiau", level=2)
    doc.add_paragraph(metadata.get("conclusion", "—"))

    doc.add_heading("I?aaei?aiey", level=2)
    doc.add_paragraph(metadata.get("suggestions", "—"))

    doc.save(output_path)
